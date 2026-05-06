from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def export_to_docx(exam: dict, include_answers: bool = False) -> bytes:
    doc = Document()

    # Title
    title = doc.add_heading(exam.get("title", "Exam Paper"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta info
    meta = doc.add_paragraph()
    meta.add_run(f"Subject: {exam.get('subject', '')}   |   "
                 f"Total Marks: {exam.get('total_marks', '')}   |   "
                 f"Duration: {exam.get('duration_minutes', '')} minutes")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("─" * 60)

    # Section A: MCQs
    doc.add_heading("SECTION A — Multiple Choice Questions", level=1)
    for i, q in enumerate(exam.get("mcqs", []), 1):
        p = doc.add_paragraph()
        p.add_run(f"Q{i}. {q['question']}").bold = True
        p.add_run(f"  [{q['bloom_level']}]").font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        for opt in q.get("options", []):
            doc.add_paragraph(f"     {opt}", style="List Bullet")
        if include_answers:
            ans = doc.add_paragraph(f"✓ Answer: {q['answer']}")
            ans.runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    # Section B: Short Answers
    doc.add_heading("SECTION B — Short Answer Questions", level=1)
    for i, q in enumerate(exam.get("short_answers", []), 1):
        p = doc.add_paragraph()
        p.add_run(f"Q{i}. {q['question']}").bold = True
        p.add_run(f"  [{q['bloom_level']}]").font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        doc.add_paragraph("Answer: ___________________________________________")
        if include_answers:
            ans = doc.add_paragraph(f"✓ {q['answer']}")
            ans.runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    # Section C: Case Studies
    doc.add_heading("SECTION C — Case Studies", level=1)
    for i, cs in enumerate(exam.get("case_studies", []), 1):
        doc.add_paragraph(f"Case Study {i}:", style="Heading 2")
        doc.add_paragraph(cs.get("scenario", ""))
        for j, (ques, ans) in enumerate(zip(cs.get("questions", []), cs.get("answers", [])), 1):
            p = doc.add_paragraph()
            p.add_run(f"  {j}. {ques}").bold = True
            if include_answers:
                a = doc.add_paragraph(f"     ✓ {ans}")
                a.runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()